import os
import re
import time
import pytesseract
import docx
import docx2txt
from PyPDF2 import PdfReader
from collections import Counter
from PIL import ImageFilter, ImageOps, Image
from langchain.schema import Document
from pdf2image import convert_from_path
from utils.redis_config import get_redis
from utils.pinecone import get_pinecone_index
from concurrent.futures import ThreadPoolExecutor
from langchain.text_splitter import RecursiveCharacterTextSplitter
from utils.file_utils import save_files, append_history, load_history
from utils.ChainManager import OptimizedChainManager, get_cached_embeddings, get_pinecone_vector_store
import uuid
from services.grpc_func import ServiceClient

clinet = ServiceClient()
chain_manager = OptimizedChainManager()
user_patterns_cache = {}

def count_words(text):
    """Count words in text"""
    return len(re.findall(r'\b\w+\b', text.strip()))

def extract_text_from_txt(path):
    """Extract text from TXT file"""
    try:
        with open(path, 'r', encoding='utf-8') as f:
            return f.read()
    except UnicodeDecodeError:
        try:
            with open(path, 'r', encoding='latin-1') as f:
                return f.read()
        except Exception as e:
            print(f"❌ TXT extraction failed for {path}: {e}")
            return ""
    except Exception as e:
        print(f"❌ TXT extraction failed for {path}: {e}")
        return ""

def extract_text_from_docx(path):
    """Extract text from DOCX file"""
    try:
        doc = docx.Document(path)
        text = ""
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        return text
    except Exception as e:
        print(f"❌ DOCX extraction failed for {path}: {e}")
        return ""

def extract_text_from_doc(path):
    """Extract text from DOC file (requires python-docx2txt)"""
    try:
        return docx2txt.process(path)
    except ImportError:
        print("❌ docx2txt not installed. Cannot process .doc files")
        return ""
    except Exception as e:
        print(f"❌ DOC extraction failed for {path}: {e}")
        return ""

def preprocess_image(img):
    """Preprocess image for better OCR"""
    return ImageOps.autocontrast(ImageOps.grayscale(img).filter(ImageFilter.SHARPEN))

def extract_text_from_image(path):
    """Extract text from image using OCR"""
    try:
        img = Image.open(path)
        processed_img = preprocess_image(img)
        text = pytesseract.image_to_string(processed_img)
        return text
    except Exception as e:
        print(f"❌ Image OCR failed for {path}: {e}")
        return ""

def extract_text_from_pdf(path):
    """Extract text from PDF using PyPDF2 and OCR fallback"""
    text = ""
    try:
        reader = PdfReader(path)
        images = convert_from_path(path)
        for i, page in enumerate(reader.pages):
            page_text = page.extract_text() or ""
            if not page_text.strip() and i < len(images):
                page_text = pytesseract.image_to_string(preprocess_image(images[i]))
            text += page_text + "\n"
    except Exception as e:
        print(f"❌ PDF extraction failed for {path}: {e}")
    return text

def extract_text_from_file(path):
    """Extract text from any supported file type"""
    file_ext = os.path.splitext(path)[1].lower()
    
    extractors = {
        '.pdf': extract_text_from_pdf,
        '.txt': extract_text_from_txt,
        '.docx': extract_text_from_docx,
        '.doc': extract_text_from_doc,
        '.png': extract_text_from_image,
        '.jpg': extract_text_from_image,
        '.jpeg': extract_text_from_image,
        '.tiff': extract_text_from_image,
        '.bmp': extract_text_from_image,
    }
    
    extractor = extractors.get(file_ext)
    if extractor:
        print(f"📄 Extracting text from {file_ext} file: {os.path.basename(path)}")
        return extractor(path)
    else:
        print(f"❌ Unsupported file type: {file_ext}")
        return ""

def extract_text_from_all_files(folder):
    """Extract text from all supported files in folder using ThreadPoolExecutor"""
    supported_extensions = {'.pdf', '.txt', '.doc', '.docx', '.png', '.jpg', '.jpeg', '.tiff', '.bmp'}
    files = []
    
    for f in os.listdir(folder):
        file_ext = '.' + f.split('.')[-1].lower()
        if file_ext in supported_extensions:
            files.append(os.path.join(folder, f))
    
    print(f"📁 Found {len(files)} supported files to process")
    
    with ThreadPoolExecutor(max_workers=4) as executor:
        texts = list(executor.map(extract_text_from_file, files))
    
    combined_text = "\n".join(filter(None, texts))
    return combined_text

def get_text_chunks(text):
    """Split text into chunks for vector storage"""
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=1000, 
        chunk_overlap=200,
        length_function=len,
        separators=["\n\n", "\n", " ", ""]
    )
    return splitter.split_text(text)

async def process_files(user_id, doc_id, files, name):
    """Process multiple file types and store in Pinecone"""
    start_time = time.time()
    
    # Try to get or create chat (with error handling)
    try:
        chat = await clinet.get_chat(doc_id)
        if not chat:
            try:
                response = await clinet.create_chat(doc_id, user_id, doc_text=name)
                if not response:
                    print(f"⚠️ Failed to create chat via gRPC, proceeding without chat creation")
                else:
                    print(f"✅ Created new chat for doc_id: {doc_id}")
            except Exception as grpc_error:
                print(f"⚠️ gRPC connection failed: {grpc_error}")
                return{
                    "error": "Chat service unavailable. Please try again later."
                }
        else:
            print(f"✅ Using existing chat for doc_id: {doc_id}")
    except Exception as e:
        print(f"⚠️ Chat service unavailable: {e}")
        print("📝 Proceeding with file processing without chat creation")
    
    # Process files
    try:
        folder = save_files(user_id, doc_id, files)
        full_text = extract_text_from_all_files(folder)
        
        if not full_text.strip():
            print(f"❌ No text extracted from {len(files)} files")
            # Log file details for debugging
            for file in files:
                print(f"📄 File: {file.filename}, Size: {file.size if hasattr(file, 'size') else 'unknown'}")
            
            return {
                "error": "No text could be extracted from the uploaded files. Please ensure your files contain readable text or try uploading different file formats (PDF, TXT, DOC, DOCX, or images with text)."
            }
        
        # Count words in extracted text
        word_count = count_words(full_text)
        print(f"📊 Extracted {word_count} words from {len(files)} files")
        
        # Check minimum word requirement
        if word_count < 20:
            return {
                "error": f"The extracted text contains only {word_count} words. A minimum of 20 words is required for processing. Please upload files with more text content."
            }
        
        print(f"✅ Text validation passed: {word_count} words extracted")
        
        await cache_document_content(doc_id, full_text)
        chunks = get_text_chunks(full_text)

        print(f"📄 Created {len(chunks)} text chunks")
        if chunks:
            print(f"📝 First chunk preview: {chunks[0][:200]}...")
            
        embeddings = await get_cached_embeddings()
        vector_store = await get_pinecone_vector_store(doc_id, embeddings)
        
        # Store in Pinecone
        index = get_pinecone_index()
        namespace = f"doc_{doc_id}"
        
        try:
            existing_docs = index.query(
                vector=[0.0] * 768, 
                namespace=namespace,
                top_k=1,
                include_metadata=True
            )
            
            if existing_docs.get('matches'):
                print(f"📝 Adding to existing document collection in Pinecone")
            else:
                print(f"📝 Creating new document collection in Pinecone")
        except Exception as e:
            print(f"⚠️ Could not check existing docs: {e}")
        
        texts_with_metadata = []
        metadatas = []
        ids = []
        
        for i, chunk in enumerate(chunks):
            if chunk.strip(): 
                chunk_id = f"{doc_id}_chunk_{i}_{uuid.uuid4().hex[:8]}"
                texts_with_metadata.append(chunk.strip())
                metadatas.append({
                    "doc_id": doc_id,
                    "chunk_id": i,
                    "user_id": user_id,
                    "doc_name": name,
                    "timestamp": time.time(),
                    "file_types": [f.filename.split('.')[-1].lower() for f in files],
                    "total_words": word_count
                })
                ids.append(chunk_id)
        
        if texts_with_metadata:
            vector_store.add_texts(
                texts=texts_with_metadata,
                metadatas=metadatas,
                ids=ids
            )
            
            processing_time = time.time() - start_time
            print(f"✅ Successfully processed {len(files)} files in {processing_time:.2f}s")
            print(f"📊 Added {len(texts_with_metadata)} chunks to Pinecone for doc_id: {doc_id}")
            
            return {
                "success": True,
                "total_words": word_count,
                "chunks_created": len(texts_with_metadata),
                "files_processed": len(files)
            }
        else:
            return {
                "error": "No valid chunks could be created from the extracted text."
            }
        
    except Exception as e:
        print(f"❌ Error in file processing: {e}")
        return {
            "error": f"Failed to process files: {str(e)}"
        }

# Keep all other functions the same (answer_question, etc.)
async def get_cached_vector_store(doc_id):
    """Get vector store from Pinecone with Redis caching"""
    redis_client = get_redis()
    cache_key = f"pinecone_store:{doc_id}"
    
    try:
        cached = await redis_client.get(cache_key)
        if cached:
            print(f"✅ Vector store metadata found in Redis for doc_id: {doc_id}")
    except Exception as e:
        print(f"⚠️ Redis cache check failed: {e}")
    
    try:
        embeddings = await get_cached_embeddings()
        vector_store = await get_pinecone_vector_store(doc_id, embeddings)
        
        try:
            await redis_client.setex(cache_key, 3600, "exists")
            print(f"✅ Cached vector store metadata for {doc_id}")
        except Exception:
            pass
        
        return vector_store
        
    except Exception as e:
        print(f"❌ Error getting vector store for {doc_id}: {e}")
        raise

async def get_cached_document_content(doc_id):
    """Get cached document content from Redis"""
    try:
        content = await get_redis().get(f"document_content:{doc_id}")
        if content:
            print(f"✅ Retrieved document content for {doc_id} from Redis")
        return content
    except Exception:
        return None

async def cache_document_content(doc_id, content):
    """Cache document content in Redis"""
    try:
        await get_redis().setex(f"document_content:{doc_id}", 7200, content)
        print(f"✅ Cached document content for {doc_id}")
    except Exception:
        pass

async def get_cached_response(doc_id, question):
    """Get cached response for frequently asked questions"""
    try:
        response = await get_redis().get(f"response:{doc_id}:{hash(question.lower().strip())}")
        if response:
            print("✅ Retrieved cached response")
        return response
    except Exception:
        return None

async def cache_response(doc_id, question, answer):
    """Cache response for future use"""
    try:
        await get_redis().setex(f"response:{doc_id}:{hash(question.lower().strip())}", 3600, answer)
        print("✅ Cached response for future use")
    except Exception:
        pass

async def learn_user_patterns(user_id, question, _):
    """Learn user patterns for better responses"""
    data = user_patterns_cache.setdefault(user_id, {
        'typo_patterns': {}, 'question_styles': [], 'preferred_length': 'medium', 'common_words': Counter()
    })
    if len(data['question_styles']) >= 10:
        data['question_styles'].pop(0)
    data['question_styles'].append(question.lower())
    data['common_words'].update(re.findall(r'\b\w+\b', question.lower()))
    if len(data['common_words']) > 50:
        data['common_words'] = Counter(dict(data['common_words'].most_common(50)))
    q = question.lower()
    if any(w in q for w in ['brief', 'short', 'quickly']):
        data['preferred_length'] = 'short'
    elif any(w in q for w in ['detail', 'explain', 'comprehensive']):
        data['preferred_length'] = 'long'

async def optimized_content_search(question, vector_store, top_k=4):
    """Optimized content search using Pinecone"""
    try:
        print(f"🔍 Searching for: '{question}' in Pinecone")
        docs = vector_store.similarity_search(question, k=top_k)
        content = []
        for doc in docs:
            if doc.page_content and doc.page_content.strip():
                content.append(doc.page_content.strip())
                print(f"📄 Found relevant content: {doc.page_content[:100]}...")
        
        print(f"✅ Found {len(content)} relevant documents from Pinecone")
        return content[:top_k]
        
    except Exception as e:
        print(f"❌ Error in content search: {e}")
        try:
            print("🔄 Trying alternative search method...")
            results = vector_store.similarity_search_with_score(question, k=top_k)
            content = [doc.page_content.strip() for doc, score in results if doc.page_content.strip()]
            print(f"✅ Alternative search found {len(content)} documents")
            return content[:top_k]
        except Exception as e2:
            print(f"❌ Alternative search also failed: {e2}")
            return []

async def answer_question(user_id, doc_id, question, is_normal_chat=False, context_only=False):
    """Answer question using unified approach with direct model calls for hybrid mode
    
    Args:
        user_id: The user ID
        doc_id: The document ID
        question: The user's question
        is_normal_chat: Whether this is a normal chat without document context
        context_only: If True, only use embedded data; if False, use Gemini directly
    """
    
    print(f"🤖 Processing question: '{question}' for doc_id: {doc_id}, normal_chat: {is_normal_chat}, context_only: {context_only}")
    cached = await get_cached_response(doc_id, question)
    if cached:
        return cached
    
    identity_keywords = ["model name", "what model", "your name", "who are you", "what are you", "ur name", "wat r u"]
    if any(k in question.lower() for k in identity_keywords):
        answer = "My name is Jack. I'm an AI assistant designed to help you understand and analyze your documents."
        await cache_response(doc_id, question, answer)
        await learn_user_patterns(user_id, question, "identity")
        await append_history(doc_id, question, answer)
        return answer
    
    try:
        # For normal chat (without document context)
        if is_normal_chat:
            if context_only:
                answer = "I can't provide a context-only answer for a normal chat as there's no document context available. Please upload a document first or ask questions about your documents."
                await cache_response(doc_id, question, answer)
                await append_history(doc_id, question, answer)
                return answer
                
            print(f"💬 Processing as normal chat for doc_id: {doc_id}")
            # Use direct model call for normal chat - clean prompt without history
            model = await chain_manager.get_direct_model(doc_id)
            
            prompt = f"""You are Jack, a helpful AI assistant. Answer the user's question directly and conversationally.

User Question: {question}

Answer:"""
            
            response = await model.ainvoke(prompt)
            answer = response.content.strip()
            
            if not answer or len(answer) < 10:
                answer = "I'm not sure I understand your question. Could you please provide more details?"
            
            await cache_response(doc_id, question, answer)
            await learn_user_patterns(user_id, question, answer)
            await append_history(doc_id, question, answer)
            
            return answer
        
        # For document-based chat
        if context_only:
            print("📄 Using context-only mode - searching embedded data only")
            vector_store = await get_cached_vector_store(doc_id)
            context = await optimized_content_search(question, vector_store, top_k=4)
            
            # Get conversation history for context-only mode
            history = load_history(doc_id)[-3:]
            history_str = "\n".join([f"Q: {h['question']}\nA: {h['answer']}" for h in history])
            
            # Get context-only chain
            chain = await chain_manager.get_chain(doc_id, context_only=True)
            
            if not context:
                docs = []
            else:
                docs = [Document(page_content=c) for c in context]
            
            response = chain({
                "input_documents": docs, 
                "question": question, 
                "history": history_str
            }, return_only_outputs=True)
            
            answer = response["output_text"].strip()
            
            if not answer or len(answer) < 10:
                answer = "The provided documents do not contain sufficient information to answer your question. Please try asking about topics that are specifically covered in your documents."
            
            await cache_response(doc_id, question, answer)
            await learn_user_patterns(user_id, question, "context_only")
            await append_history(doc_id, question, answer)
            return answer
        
        else:
            print("🔄 Using hybrid mode - Pure Gemini response without document interference")
            
            # Get direct model for hybrid mode
            model = await chain_manager.get_direct_model(doc_id)
            
            # Create clean prompt for pure Gemini response without context or history interference
            prompt = f"""You are Jack, a helpful AI assistant. Answer the user's question directly using your knowledge.

Instructions:
- Provide a clear, direct answer to the question
- Use your general knowledge to give the best response
- Be conversational and helpful
- Don't reference any documents or previous conversations

User Question: {question}

Answer:"""
            
            response = await model.ainvoke(prompt)
            answer = response.content.strip()
            
            if not answer or len(answer) < 10:
                answer = "I'm having trouble generating a response. Could you please try rephrasing your question?"
            
            await cache_response(doc_id, question, answer)
            await learn_user_patterns(user_id, question, answer)
            await append_history(doc_id, question, answer)
            
            return answer
        
    except Exception as e:
        print(f"❌ Error in answer_question: {e}")
        fallback = "I apologize, but I encountered an issue processing your question. Could you please try rephrasing it?"
        await append_history(doc_id, question, fallback)
        return fallback
